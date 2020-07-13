import os
import re

import docx
import unicodedata
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Cm
from nltk import word_tokenize, sent_tokenize
from tika import parser
from docx import Document

from src.python.helper import archivos_entrenamiento_limpios
from src.python.metodos_de_similitud import obtener_similitud


class ArchivoTxt:
    def __init__(self, nombre, extension, texto):
        self.nombre = nombre
        self.extension = extension
        self.texto = texto

def obtener_archivos(nombre_directorio):
    archivos = os.listdir(nombre_directorio)
    return [convertir_archivo_a_txt(nombre_directorio, archivo) for archivo in archivos]

def convertir_documento_a_txt(archivo, nombre_directorio):
    raw = parser.from_file(nombre_directorio + archivo)
    return raw['content']

def convertir_archivo_a_txt(nombre_directorio, archivo):
    archivo_nombre, archivo_extension = os.path.splitext(archivo)
    archivo_txt = convertir_documento_a_txt(archivo, nombre_directorio)
    return ArchivoTxt(archivo_nombre, archivo_extension, archivo_txt)

def limpiar(archivo):
    archivo_limpio = re.sub(r'\n+', '\n', archivo.strip()) # reemplazo multiples enter por uno solo
    archivo_limpio = re.sub('\n', '. ', archivo_limpio.strip())
    archivo_limpio = re.sub(r'[.][.]+', '.', archivo_limpio.strip())
    archivo_limpio = re.sub(r'[ ][ ]+', ' ', archivo_limpio.strip())
    archivo_limpio = re.sub('á', 'a', archivo_limpio.strip())
    archivo_limpio = re.sub('é', 'e', archivo_limpio.strip())
    archivo_limpio = re.sub('í', 'i', archivo_limpio.strip())
    archivo_limpio = re.sub('ó', 'o', archivo_limpio.strip())
    archivo_limpio = re.sub('ú', 'u', archivo_limpio.strip())
    archivo_limpio = re.sub('”', '"', archivo_limpio.strip())
    archivo_limpio = re.sub('“', '"', archivo_limpio.strip())
    archivo_limpio = re.sub('\u200b', ' ', archivo_limpio.strip())

    archivo_limpio = unicodedata.normalize("NFKD", archivo_limpio.strip())

    oraciones = sent_tokenize(archivo_limpio.strip(), "spanish")
    oraciones_limpias = []
    for oracion in oraciones:
        if oracion.strip() != '.':
            if oracion.strip().endswith('.'):
                oracion_a_agregar = oracion[:-1]
            else:
                oracion_a_agregar = oracion
            oraciones_limpias.append(oracion_a_agregar.strip())

    i=0
    j=0
    # TODO: para arreglar enters que deberian ser espacios para que siga la oracion (pasa en pdfs nomas)
    oraciones_mas_limpias = []
    while i < len(oraciones_limpias):
        if i == 0:
            oraciones_mas_limpias.append(oraciones_limpias[0])
        else:
            palabras_oracion = word_tokenize(oraciones_limpias[i])
            if palabras_oracion[0].islower():
                oraciones_mas_limpias[j] += " " + oraciones_limpias[i]
            else:
                j += 1
                oraciones_mas_limpias.append(oraciones_limpias[i])
        i += 1

    return oraciones_mas_limpias

def limpiar_archivos_entrenamiento(archivo):
    archivo_limpio = limpiar(archivo.texto)
    archivos_entrenamiento_limpios.append(ArchivoTxt(archivo.nombre, archivo.extension, archivo_limpio))

def correctamente_citada(url, texto_archivo_test_limpio):
    for oracion in texto_archivo_test_limpio:
        similitud = obtener_similitud(oracion, url)
        if similitud >= 0.99:
            return True
    return False


def excluida(oracion):
    archivo_text_excluido = convertir_archivo_a_txt("../../", "Texto excluido de plagio.txt")
    if archivo_text_excluido.texto is None:
        return False
    archivo_limpio = limpiar(archivo_text_excluido.texto)
    for oracion_archivo in archivo_limpio:
        similitud = obtener_similitud(oracion, oracion_archivo)
        if similitud > 0.9:
            return True
    return False

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def guardar_resultado(nombre_archivo, nombre_alumno, topico_con_mas_score, plagio, tiempo_que_tardo, porcentaje_de_plagio, path_resultado, path_entrenamiento):
    document = Document()

    h = document.add_heading(f'Análisis de plagio sobre:\n', 0)
    h.add_run(nombre_archivo).italic = True

    p = document.add_paragraph('Tópicos del texto: ')
    p.add_run(", ".join(["a","b"])).italic = True

    if nombre_alumno:
        p = document.add_paragraph('Nombre del alumno que realizó el TP: ')
        p.add_run(nombre_alumno[0]).italic = True

    document.add_heading('Análisis de plagio', level=1)
    document.add_paragraph(f'Total de {len(plagio)} plagios encontrados en {tiempo_que_tardo}')
    document.add_paragraph(f'Porcentaje de plagio general: {porcentaje_de_plagio}%')

    table = document.add_table(rows=1, cols=4)
    table.allow_autofit = False
    table.style = 'Medium Shading 1 Accent 1' # Nombres de estilos de tablas en word
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Oración plagiada'
    hdr_cells[1].text = 'Oración original'
    hdr_cells[2].text = 'Lugar donde se encontró'
    hdr_cells[3].text = 'Ubicación'

    for oracion, plagio, porcentaje, url, ubicacion in plagio:
        row_cells = table.add_row().cells
        row_cells[0].text = oracion
        row_cells[1].text = plagio
        p = row_cells[2].add_paragraph()
        if str(url).startswith("http"):
            add_hyperlink(p, url, url)
        else:
            add_hyperlink(p, url, os.path.abspath(path_entrenamiento) + '\\' + url)
        row_cells[3].text = ubicacion

    nombre_archivo_plagio = path_resultado + 'Plagio ' + str(str(nombre_archivo).split(".")[0]) + '.docx'
    document.save(nombre_archivo_plagio)
